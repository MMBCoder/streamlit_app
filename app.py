import streamlit as st
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
import docx
import os

# For local development, attempt to load from .env if not on Streamlit Cloud
try:
    openai_api_key = st.secrets["OPENAI_API_KEY"]
except Exception:
    from dotenv import load_dotenv
    load_dotenv()
    openai_api_key = os.getenv("OPENAI_API_KEY")

# Verify the API key is loaded
if not openai_api_key:
    st.error("OpenAI API key not found! Please set it in your Streamlit secrets or your .env file.")
    st.stop()

# Load and process SAS code from a Word file
def load_sas_code_from_word(file_path):
    doc = docx.Document(file_path)
    sas_code = [para.text for para in doc.paragraphs]
    return '\n'.join(sas_code)

# Load campaign requirements from a Word file
def load_campaign_requirements(file_path):
    doc = docx.Document(file_path)
    requirements = [para.text for para in doc.paragraphs]
    return '\n'.join(requirements)

# Streamlit UI
st.title("AI-Based Campaign Operation Programming Solution")

# Load files using relative paths (assuming both doc files are in the same directory as app.py)
sas_code_path = "sas_code_example.docx"
requirements_path = "campaign_requirements_example.docx"

if os.path.exists(sas_code_path) and os.path.exists(requirements_path):
    sas_code_text = load_sas_code_from_word(sas_code_path)
    requirements_text = load_campaign_requirements(requirements_path)

    st.success("SAS code and Campaign requirements loaded Successfully in Vector DB")

    # Chunking the SAS code into smaller pieces
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50,
        separators=["\n\n", "\n", ".", " "]
    )
    sas_code_chunks = text_splitter.split_text(sas_code_text)

    # Creating FAISS vector store
    embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
    vector_db = FAISS.from_texts(sas_code_chunks, embeddings)

    # User input
    user_query = st.text_area("Enter Your Campaign Requirement")

    if st.button("Submit") and user_query:
        llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0, openai_api_key=openai_api_key)

        prompt_template = PromptTemplate(
            input_variables=["context", "question"],
            template=(
                "You are a SAS programming expert. Based on the provided SAS code snippets below, "
                "generate SAS code that fulfills the user's campaign requirements.\n\n"
                "Relevant SAS Code Snippets:\n{context}\n\n"
                "User's Campaign Requirement:\n{question}\n\n"
                "Provide the full SAS code that accomplishes this task."
            ),
        )

        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=vector_db.as_retriever(),
            chain_type_kwargs={"prompt": prompt_template},
        )

        response = qa_chain.run(user_query)

        st.write("### Generated SAS Code:")
        st.code(response, language='sas')

else:
    st.error("Campaign Requirements Data not found in COSMOS DB")

# Sidebar with Image and Instructions
st.sidebar.image("https://www.streamlit.io/images/brand/streamlit-mark-color.png", width=150)
st.sidebar.title("Instructions")
st.sidebar.write("1. Install required libraries:")
st.sidebar.code("pip install streamlit langchain faiss-cpu python-docx openai tiktoken")
st.sidebar.write("2. Run the Streamlit app:")
st.sidebar.code("streamlit run app.py")
st.sidebar.write("3. If you encounter 'ModuleNotFoundError: No module named micropip', try creating a virtual environment:")
st.sidebar.code(
    "python -m venv venv\n"
    "venv\\Scripts\\activate (Windows)\n"
    "source venv/bin/activate (Linux/Mac)\n"
    "pip install streamlit langchain faiss-cpu python-docx openai tiktoken"
)